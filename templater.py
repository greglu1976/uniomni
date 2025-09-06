# ГЕНЕРАЦИЯ И ЗАПОЛНЕНИЕ ШАБЛОНА БЛАНКА УСТАВОК
import os

from docxtpl import DocxTemplate
from docx import Document

from docx_handler import add_new_section, add_new_section_landscape
from tables import add_table_settings, add_table_reg, add_table_fks, add_table_leds_new, add_table_mtrx_ins, add_table_mtrx_outs, add_table_binaries, add_table_final

def fill_template(fsu, hardware):

    doc = DocxTemplate("temp.docx")

    # Формируем контекст для Jinja2
    context = {
        "fsu": fsu,
        "hardware": hardware,
    }
 
    # Заполняем документ
    doc.render(context)
    doc.save("Бланк уставок.docx")


def _create_config_disturb(doc):

    p = doc.add_paragraph('Настройка регистрации')
    p.style = 'ДОК Заголовок 2'

    p = doc.add_paragraph(r'Резервирование{% set items = hardware.get_config_disturb()["Общие параметры"] %}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)

    p = doc.add_paragraph()
    p.style = 'TAGS'

def _create_config_cpu(doc):

    p = doc.add_paragraph('Модуль ЦП')
    p.style = 'ДОК Заголовок 2'

    p = doc.add_paragraph(r'Резервирование{% set items = hardware.get_config_cpu()["Резервирование"] %}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)

    p = doc.add_paragraph()
    p.style = 'TAGS'

def _create_config_sync(doc):

    p = doc.add_paragraph('Синхронизация времени')
    p.style = 'ДОК Заголовок 2'

    p = doc.add_paragraph(r'Общие настройки синхронизации{% set items = hardware.get_config_sync()["Общие параметры"] %}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)

    p = doc.add_paragraph()
    p.style = 'TAGS'

    p = doc.add_paragraph(r'Параметры летнего времени{% set items = hardware.get_config_sync()["Параметры летнего времени"] %}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)

    p = doc.add_paragraph()
    p.style = 'TAGS'

    p = doc.add_paragraph(r'Параметры зимнего времени{% set items = hardware.get_config_sync()["Параметры зимнего времени"] %}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)

def _create_config_modules(doc):
     ############################################################################
    # СОЗДАЕМ РАЗДЕЛ С КОНФИГУРАЦИЕЙ ВХОДОВ ВЫХОДОВ

    #p = doc.add_paragraph('Модули'+r'{% for plate in hardware.get_hw_plates() if hardware.get_hw_plates() %}')
    #p.style = 'ДОК Заголовок 2'

    p = doc.add_paragraph(r'{% for plate in hardware.get_hw_plates() if hardware.get_hw_plates() %}')
    p.style = 'TAGS'

    p = doc.add_paragraph(r'Слот М{{ plate.get_slot() }}. {{ plate.get_name() }}'+ r'{% set items = plate.get_info() %}' + r'{% if items %}')
    #p.style = 'ДОК Заголовок 3'
    p.style = 'ДОК Заголовок 2'

    p = doc.add_paragraph(r'Общие настройки конфигурации')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)
    p = doc.add_paragraph(r'{% endif %}'+r'{% for dics in plate.get_inputs() if plate.get_inputs() %}' + r'{% set items = dics.data %}')
    p.style = 'TAGS'

    p1 = doc.add_paragraph(r'Дискретный вход {{ dics.counter }}')
    p1.style = 'ДОК Таблица Название'
    add_table_binaries(doc)
    p = doc.add_paragraph(r'{% endfor %}'+r'{% for dics in plate.get_outputs() if plate.get_outputs() %}'  + r'{% set items = dics.data %}')
    p.style = 'TAGS'

    p = doc.add_paragraph(r'Выходное реле {{ dics.counter }}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)
    p = doc.add_paragraph(r'{% endfor %}'+r'{% for dics in plate.get_volts() if plate.get_volts() %}' + r'{% set items = dics.data %}')
    p.style = 'TAGS'

    p = doc.add_paragraph(r'Измерительный вход напряжения {{ dics.counter }}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)
    #p = doc.add_paragraph(r'{% endfor %}'+r'{% for items in plate.get_currents() if plate.get_currents() %}')
    p = doc.add_paragraph(r'{% endfor %}'+r'{% for dics in plate.get_currents() if plate.get_currents() %}' + r'{% set items = dics.data %}')
    p.style = 'TAGS'

    p = doc.add_paragraph(r'Измерительный токовый вход {{ dics.counter }}')
    p.style = 'ДОК Таблица Название'
    add_table_binaries(doc)
    p = doc.add_paragraph(r'{% endfor %}{% endfor %}')
    p.style = 'TAGS'

# СУММАРНАЯ ФУНКЦИЯ СОЗДАНИЯ РАЗДЕЛА КОНФИГУРАЦИЯ
def _create_section_config(doc):

    ############################################################################
    # СОЗДАЕМ РАЗДЕЛ С КОНФИГУРАЦИЕЙ
    add_new_section(doc)
    p = doc.add_paragraph('КОНФИГУРАЦИЯ')
    p.style = 'ДОК Заголовок 1'

    # конфигурация синхронизация времени
    _create_config_sync(doc)
    # конфигурация ЦП
    _create_config_cpu(doc)
    # Настройка регистрации
    _create_config_disturb(doc)
    # конфигурация модулей
    _create_config_modules(doc)

# СУММАРНАЯ ФУНКЦИЯ СОЗДАНИЯ РАЗДЕЛА УСТАВКИ РЗиА
def _create_section_settings(doc):
    add_new_section_landscape(doc)

    p = doc.add_paragraph('УСТАВКИ РЗиА'+r'{% for fb in fsu.get_fbs() if fb.is_fb_settings_empty() %}')
    p.style = 'ДОК Заголовок 1'

    p = doc.add_paragraph(r'{{ fb.get_description() }} ({{ fb.get_fb_name() }}) {% for func in fb.get_functions() if func.get_settings_for_bu() %}')
    #p.style = 'ДОК Заголовок 3'
    p.style = 'ДОК Заголовок 2'

    p = doc.add_paragraph(r'{{ func.get_description() }}{% if func.get_name() %} ({{ func.get_name() }}){% endif %}')
    p.style = 'ДОК Таблица Название'

    add_table_settings(doc)

    p = doc.add_paragraph(r'{% endfor %}{% endfor %}')
    p.style = 'TAGS'
# РАЗДЕЛ ПАРАМЕТРИРОВАНИЯ ВХОДОВ И ВЫХОДОВ
def _create_section_inouts(fsu, doc):

    add_new_section_landscape(doc) # Создаем раздел для матрицы вх/вых
    # Добавляем заголовок
    p = doc.add_paragraph('МАТРИЦА ВХОДОВ И ВЫХОДНЫХ РЕЛЕ')
    p.style = 'ДОК Заголовок 1'

    p = doc.add_paragraph('Дискретные входы')
    p.style = 'ДОК Заголовок 2'

    text = doc.add_paragraph('Для дискретного входа возможно подключение только одного сигнала.')
    text.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{% for input_plate in hardware.get_inputs() if hardware.get_inputs() %}'+r'{% if input_plate.num_of_inputs %}')
    p.style = 'ДОК Текст'
    
    p = doc.add_paragraph(r'{{ input_plate.desc }}')
    p.style = 'ДОК Таблица Название'

    inputs_ = fsu.get_fsu_inputs_list()
    inputs = set([item['Полное наименование сигнала'] for item in inputs_])
    inputs = sorted(list(inputs))

    control_ = fsu.get_fsu_control_list()
    controls = set([item['Полное наименование сигнала'] for item in control_])
    controls = sorted(list(controls))

    add_table_mtrx_ins(doc, inputs, controls)

    p = doc.add_paragraph(r'{% endif %}{% endfor %}')
    p.style = 'TAGS'   

    ############################### ВЫХОДНЫЕ РЕЛЕ ####################################

    p = doc.add_paragraph('Выходные реле')
    p.style = 'ДОК Заголовок 2'

    text = doc.add_paragraph('Возможно подключение до пяти сигналов на одно выходное реле.')
    text.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{% for output_plate in hardware.get_outputs() if hardware.get_outputs() %}'+r'{% if output_plate.num_of_outputs %}')
    p.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{{ output_plate.desc }}')
    p.style = 'ДОК Таблица Название'

    statuses_ = fsu.get_fsu_statuses()
    statuses = set([item['Полное наименование сигнала'] for item in statuses_])
    statuses = sorted(list(statuses))

    add_table_mtrx_outs(doc, statuses)

    p = doc.add_paragraph(r'{% endif %}{% endfor %}')
    p.style = 'TAGS' 

# РАЗДЕЛ СВЕТОДИОДОВ И ФК
def _create_section_leds(fsu, doc):
    #############################################################################
    # СОЗДАЕМ РАЗДЕЛ НАСТРОЙКА СВЕТОДИОДОВ И ФУНКЦИОНАЛЬНЫХ КЛАВИШ
    add_new_section_landscape(doc) # Создаем раздел для матрицы вх/вых
    # Добавляем заголовок
    p = doc.add_paragraph('НАСТРОЙКА СВЕТОДИОДОВ И ФУНКЦИОНАЛЬНЫХ КЛАВИШ')
    p.style = 'ДОК Заголовок 1'

    ###############################################################
    p = doc.add_paragraph('Светодиоды')
    p.style = 'ДОК Заголовок 2'

    text = doc.add_paragraph('Для светодиода возможно подключение до пяти сигналов.')
    text.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{% for leds_unit in hardware.get_hmi_leds() if hardware.get_hmi_leds() %}')
    p.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{{ leds_unit }}')
    p.style = 'ДОК Таблица Название'

    statuses_ = fsu.get_fsu_statuses()
    statuses = set([item['Полное наименование сигнала'] for item in statuses_])
    statuses = sorted(list(statuses))

    add_table_leds_new(doc, statuses)

    p = doc.add_paragraph(r'{% endfor %}')
    p.style = 'TAGS'    

    ###############################################################
    p = doc.add_paragraph('Функциональные клавиши')
    p.style = 'ДОК Заголовок 2'

    text = doc.add_paragraph('Для функциональной клавиши возможно подключение только одного управляющего сигнала.')
    text.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{% for fks_unit in hardware.get_hmi_fks() if hardware.get_hmi_fks() %}')
    p.style = 'ДОК Текст'

    p = doc.add_paragraph(r'{{ fks_unit }}')
    p.style = 'ДОК Таблица Название'

    choices_ = fsu.get_fsu_control_list()
    choices = set([item['Наименование сигналов на ФСУ'] for item in choices_])
    choices = sorted(list(choices))

    add_table_fks(doc, choices)

    p = doc.add_paragraph(r'{% endfor %}')
    p.style = 'TAGS'

# РАЗДЕЛ НАСТРОЙКИ ПАРАМЕТРОВ РЕГИСТРАЦИИ
def _create_section_disturb(fsu, doc):
    
    add_new_section_landscape(doc)

    p = doc.add_paragraph('НАСТРОЙКА ПАРАМЕТРОВ РЕГИСТРАЦИИ')
    p.style = 'ДОК Заголовок 1'

    text = doc.add_paragraph('Возможна регистрация не более 200 сигналов.')
    text.style = 'ДОК Текст'

    # Таблица Выходные сигналы общей логики
    
    if fsu.get_fsu_sys_statuses_sorted():
        p = doc.add_paragraph('Системные сигналы')
        p.style = 'ДОК Таблица Название'
        add_table_reg(doc, generate = 0)

    # Таблица Выходные сигналы общей логики
    p = doc.add_paragraph('Выходные сигналы общей логики')
    p.style = 'ДОК Таблица Название'
    add_table_reg(doc, generate = 1) # 0 - сист сигналы, 1 - статусы, 2 - вирт ключи и кнопки, 3 - входные сигналы

    # Таблица Виртуальные ключи и клавиши
    p = doc.add_paragraph('Виртуальные ключи и клавиши')
    p.style = 'ДОК Таблица Название'
    add_table_reg(doc, generate = 2)

    # Таблица Входные дискретные сигналы
    p = doc.add_paragraph('Входные дискретные сигналы')
    p.style = 'ДОК Таблица Название'
    add_table_reg(doc, generate = 3)
########################################################################################################################################################
############################################## ГЛАВНАЯ ФУНКЦИЯ СОЗДАНИЯ ШАБЛОНА DOCX ###################################################################
def create_template(fsu, hardware):

    doc = Document('origin.docx')

    # СОЗДАЕМ РАЗДЕЛ С УСТАВКАМИ РЗИА
    _create_section_settings(doc)

    # СОЗДАЕМ РАЗДЕЛ С ПАРАМЕТРИРОВАНИЕ ДИСКРЕТНЫХ ВХОДОВ И ВЫХОДНЫХ РЕЛЕ
    _create_section_inouts(fsu, doc)

    # СОЗДАЕМ РАЗДЕЛ С ПАРАМЕТРИРОВАНИЕ СВЕТОДИОДОВ И ФК
    _create_section_leds(fsu, doc) 

    #for plate in hardware.get_hw_plates():
        #print(plate.get_name())


    # СОЗДАЕМ РАЗДЕЛ С КОНФИГУРАЦИЕЙ
    _create_section_config(doc)

    # СОЗДАЕМ РАЗДЕЛ С ПАРАМЕТРАМИ РЕГИСТРАЦИИ
    _create_section_disturb(fsu, doc)

    # ДОБАВЛЯЕМ ФИНАЛЬНУЮ ТАБЛИЦУ С ПОДПИСЯМИ
    add_table_final(doc)

    doc.save("temp.docx")
