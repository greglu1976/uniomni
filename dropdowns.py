from docx.oxml import parse_xml
from docx.oxml.shared import qn
from docx import Document

def add_formatted_dropdown(paragraph, choices, default="", alias="", 
                         label="", style=None, instruction_text="Выберите значение"):
    """
    Добавляет форматированный выпадающий список с меткой
    
    Args:
        paragraph: параграф для добавления
        choices: список вариантов
        default: значение по умолчанию
        alias: уникальное имя
        label: текст метки перед списком
        style: стиль форматирования
        instruction_text: текст подсказки
    """
    # Добавляем метку, если она указана
    if label:
        run = paragraph.add_run(f"{label}: ")
        if style:
            run.style = style

    # Создаем XML для выпадающего списка
    sdt = parse_xml(f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{alias}"/>
                <w:tag w:val="{alias}"/>
                <w:id w:val="{abs(hash(alias))}"/>
                <w:dropDownList>
                    {''.join(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>' for choice in choices)}
                </w:dropDownList>
                <w:placeholder>
                    <w:docPart w:val="{instruction_text}"/>
                </w:placeholder>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="auto"/>
                        <w:sz w:val="24"/>
                    </w:rPr>
                    <w:t>{default if default else choices[0]}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    ''')
    
    paragraph._p.append(sdt)


def add_formatted_dropdown2(paragraph, choices, default="Не назначено", alias="", instruction_text=""):
    sdt = parse_xml(f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{alias}"/>
                <w:tag w:val="{alias}"/>
                <w:id w:val="{abs(hash(alias))}"/>
                <w:dropDownList>
                    <w:listItem w:displayText="{default}" w:value=""/>
                    {''.join(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>' for choice in choices)}
                </w:dropDownList>
                <w:showingPlcHdr/>
                <w:placeholder>
                    <w:docPart w:val="{instruction_text}"/>
                </w:placeholder>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="808080"/>
                        <w:sz w:val="24"/>
                        <w:i/> <!-- курсив -->
                        <!-- <w:b/> полужирный -->
                        <w:spacing w:val="10"/> <!-- интервал между символами -->
                    </w:rPr>
                    <w:t>{default}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    ''')
    
    paragraph._p.append(sdt)



def add_formatted_dropdown3(paragraph, inputs_choices, controls_choices=[], default="Не назначено", alias="", instruction_text=""):
    # Формируем элементы списка с разделителями
    list_items = []
    
    # Добавляем inputs
    if inputs_choices:
        list_items.append('<w:listItem w:displayText="──────── Сигналы РЗиА ────────" w:value="INPUTS_HEADER" w:disabled="true"/>')
        for choice in inputs_choices:
            list_items.append(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>')
    
    # Добавляем controls
    if controls_choices:
        if inputs_choices:  # Добавляем разделитель только если есть оба списка
            list_items.append('<w:listItem w:displayText=" " w:value="SPACER" w:disabled="true"/>')
        list_items.append('<w:listItem w:displayText="────── Общие сигналы ФС ──────" w:value="CONTROLS_HEADER" w:disabled="true"/>')
        for choice in controls_choices:
            list_items.append(f'<w:listItem w:displayText="{choice}" w:value="{choice}"/>')
    
    sdt = parse_xml(f'''
        <w:sdt xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:sdtPr>
                <w:alias w:val="{alias}"/>
                <w:tag w:val="{alias}"/>
                <w:id w:val="{abs(hash(alias))}"/>
                <w:dropDownList>
                    <w:listItem w:displayText="{default}" w:value=""/>
                    {''.join(list_items)}
                </w:dropDownList>
                <w:showingPlcHdr/>
                <w:placeholder>
                    <w:docPart w:val="{instruction_text}"/>
                </w:placeholder>
            </w:sdtPr>
            <w:sdtContent>
                <w:r>
                    <w:rPr>
                        <w:color w:val="808080"/>
                        <w:sz w:val="24"/>
                        <w:i/>
                        <w:spacing w:val="10"/>
                    </w:rPr>
                    <w:t>{default}</w:t>
                </w:r>
            </w:sdtContent>
        </w:sdt>
    ''')
    
    paragraph._p.append(sdt)











if __name__ == '__main__':
    # Пример использования расширенной версии:
    doc = Document()

    # Добавляем выпадающий список с меткой и форматированием
    paragraph = doc.add_paragraph()
    choices = ["Высокий", "Средний", "Низкий"]
    add_formatted_dropdown(
        paragraph=paragraph,
        choices=choices,
        default="Выберите уровень",
        alias="PriorityLevel",
        label="Приоритет",
        instruction_text="Выберите уровень приоритета"
    )

    # Добавляем несколько выпадающих списков в таблицу
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'

    # Заполняем таблицу выпадающими списками
    for i, (label, choices) in enumerate([
        ("Статус", ["Активен", "Неактивен", "В обработке"]),
        ("Категория", ["A", "B", "C"]),
        ("Важность", ["Критическая", "Высокая", "Средняя", "Низкая"])
    ]):
        cell = table.cell(i, 0)
        cell.text = label
        
        cell = table.cell(i, 1)
        paragraph = cell.paragraphs[0]
        add_formatted_dropdown(
            paragraph=paragraph,
            choices=choices,
            #alias= f"DropDown_{i}",
            instruction_text=f"Выберите {label.lower()}",
        )

    doc.save("formatted_dropdowns.docx")
