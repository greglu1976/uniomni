# вынесено сюда всякое создание разделов 

from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.shared import Inches

def add_new_section(doc):
    """Добавляет портретную секцию"""
    if len(doc.sections) == 0:
        return doc
    
    current_section = doc.sections[-1]
    
    # Добавляем разрыв страницы перед созданием новой секции
    doc.add_paragraph()
    
    # Создаем новую секцию
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Копируем основные настройки
    new_section.orientation = WD_ORIENTATION.PORTRAIT
    new_section.page_width = Inches(8.3)  # Стандартная ширина для портрета
    new_section.page_height = Inches(11.7)  # Стандартная высота для портрета
    
    # Копируем поля
    new_section.left_margin = current_section.left_margin
    new_section.right_margin = current_section.right_margin
    new_section.top_margin = current_section.top_margin
    new_section.bottom_margin = current_section.bottom_margin
    
    return doc

def add_new_section_landscape(doc):
    """Добавляет альбомную секцию"""
    if len(doc.sections) == 0:
        return doc
    
    current_section = doc.sections[-1]
    
    # Добавляем разрыв страницы перед созданием новой секции
    doc.add_paragraph()
    
    # Создаем новую секцию
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    
    # Устанавливаем альбомную ориентацию
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width = Inches(11.7)   # Ширина становится высотой
    new_section.page_height = Inches(8.3) # Высота становится шириной
    
    # Копируем поля
    #new_section.left_margin = current_section.left_margin
    #new_section.right_margin = current_section.right_margin
    #new_section.top_margin = current_section.top_margin
    #new_section.bottom_margin = current_section.bottom_margin

    new_section.left_margin = current_section.top_margin
    new_section.right_margin = current_section.bottom_margin
    new_section.top_margin = current_section.right_margin
    new_section.bottom_margin = current_section.left_margin

    return doc

def add_new_section_test(doc):
    # Получаем текущий раздел
    current_section = doc.sections[-1]

    # Добавляем новый раздел
    new_section = doc.add_section()

    # Устанавливаем альбомную ориентацию
    new_section.orientation = WD_ORIENTATION.LANDSCAPE
    new_section.page_width = current_section.page_height
    new_section.page_height = current_section.page_width

    # Отключаем связь с предыдущими колонтитулами
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False

    # Создаем новые колонтитулы для альбомной ориентации
    header = new_section.header
    footer = new_section.footer

    # Добавляем содержимое в колонтитулы
    header_para = header.paragraphs[0]
    footer_para = footer.paragraphs[0]

    # Здесь добавьте необходимое содержимое колонтитулов
    header_para.text = "Ваш верхний колонтитул"
    footer_para.text = "Ваш нижний колонтитул"

    # Устанавливаем стиль и форматирование для колонтитулов
    header_para.style = 'Header'  # или ваш собственный стиль
    footer_para.style = 'Footer'  # или ваш собственный стиль


