from docx.shared import Cm, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
import docx
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import os, sys
import json

from pathlib import Path

from dropdowns import add_formatted_dropdown2

def set_vertical_cell_direction(cell: _Cell, direction: str):
    # direction: tbRl -- top to bottom, btLr -- bottom to top
    assert direction in ("tbRl", "btLr")
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)  # btLr tbRl
    tcPr.append(textDirection)

def set_repeat_table_header(row):
    """ set repeat table row on every new page
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def set_cell_vertical_alignment(cell, align="center"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcValign = OxmlElement('w:vAlign')
        tcValign.set(qn('w:val'), align)
        tcPr.append(tcValign)

def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


####################################################################################
################################ ТАБЛИЦА ДЛЯ УСТАВОК ##############################
####################################################################################

table_settings = (Inches(0.28), Inches(1.23), Inches(0.9), Inches(0.5), Inches(1.5), Inches(0.55), Inches(0.45), Inches(0.9), Inches(1.05))  #задаем ширину столбцов таблицы вывода репортов

def add_table_settings(doc):
    table = doc.add_table(rows=5, cols=9)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Описание'
    hdr_cells[2].text = 'Наименование'
    hdr_cells[4].text = 'Значение / Диапазон'
    hdr_cells[5].text = 'Ед. изм.'
    hdr_cells[6].text = 'Шаг'
    hdr_cells[7].text = 'Значение по умолчанию'
    hdr_cells[8].text = 'Уставка'
    for i in range(0,9):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.runs[0].font.size = Pt(10)

    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[2].text = 'ПО'
    hdr_cells[3].text = 'ФСУ'
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # третья строка со служебными тегами
    hdr_cells = table.rows[2].cells
    #hdr_cells[2].text = '{%tr for param_name, param_data in input_value.properties.items() %}'
    tag = f'for row in func.get_settings_for_bu()'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'
    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{{ loop.index }}'
    hdr_cells[1].text = '{{ row["Описание"] }}'
    hdr_cells[2].text = '{{ row["Наименование ПО"] }}'
    hdr_cells[3].text = '{{ row["Наименование ФСУ"] }}'    
    hdr_cells[4].text = '{{ row["Значение / Диапазон"]  }}'
    hdr_cells[5].text = '{{ row["Ед.изм."] }}'
    hdr_cells[6].text = '{{ row["Шаг"] }}'
    hdr_cells[7].text = '{{ row["Значение по умолчанию"] }}'
    hdr_cells[8].text = '' #'{{ param_data.setpoint }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,9):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(0, 6).merge(table.cell(1, 6))
    table.cell(0, 7).merge(table.cell(1, 7))
    table.cell(0, 8).merge(table.cell(1, 8))

    table.cell(2, 0).merge(table.cell(2, 8))
    table.cell(4, 0).merge(table.cell(4, 8))

    for row in table.rows:
        for idx, width in enumerate(table_settings):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить

        # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Устанавливаем размер шрифта 12 пунктов

    return table

####################################################################################
################################ КОНЕЦ ТАБЛИЦА ДЛЯ УСТАВОК #########################
####################################################################################

####################################################################################
################################ ТАБЛИЦА ДЛЯ РЕГИСТРАЦИИ ###########################
####################################################################################


table_reg = (Inches(4.5), Inches(1.5), Inches(1.6), Inches(1.6), Inches(1.6))  #задаем ширину столбцов таблицы вывода репортов

def add_table_reg(doc): # новая таблица исходящих отчетов
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр'
    hdr_cells[2].text = 'Журнал событий регистрация'
    hdr_cells[3].text = 'Осциллограф пуск'
    hdr_cells[4].text = 'Осциллограф регистрация'
    for i in range(0,5):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице


    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Обозначение ФСУ'
    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    hdr_cells = table.rows[2].cells
    tag = f'for row in fsu.get_fsu_statuses()'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{{ row["Полное наименование сигнала"] }}'
    hdr_cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'

    #hdr_cells[2].text = '{{ param_data.log }}'
    choices_start = ["Не выполняется", "По переднему фронту", "По заднему фронту", "По любому изменению"]
    par3 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par3,
        choices=choices_start,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #hdr_cells[3].text = '{{ param_data.oscill_start }}'
    par2 = hdr_cells[3].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=choices_start,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #hdr_cells[4].text = '+'
    choices_reg = ["Выведено", "Введено"]
    par1 = hdr_cells[4].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=choices_reg,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,5):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(1, 3))
    table.cell(0, 4).merge(table.cell(1, 4))

    table.cell(2, 0).merge(table.cell(2, 4))
    table.cell(4, 0).merge(table.cell(4, 4))

    for row in table.rows:
        for idx, width in enumerate(table_reg):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить

            # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Устанавливаем размер шрифта 12 пунктов

    return table    

####################################################################################
################################ КОНЕЦ ТАБЛИЦА ДЛЯ РЕГИСТРАЦИИ #####################
####################################################################################


####################################################################################
############################ ТАБЛИЦА ДЛЯ МАТРИЦЫ ВЫХОДНЫХ РЕЛЕ #####################
####################################################################################

table_mtrx_outs = (Inches(2), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7), Inches(1.7))

def add_table_mtrx_outs(doc, statuses): # новая таблица исходящих отчетов
    table = doc.add_table(rows=5, cols=6)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Выходное реле'
    hdr_cells[1].text = 'Назначенные сигналы'

    for i in range(0,6):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[1].text = '1'
    hdr_cells[2].text = '2'
    hdr_cells[3].text = '3'
    hdr_cells[4].text = '4'
    hdr_cells[5].text = '5'    
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    hdr_cells = table.rows[2].cells
    tag = r'for i in range(1, output_plate.num_of_outputs|int+1)'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = 'Реле '+'{{ loop.index }}'


    #choices_start = ["Не выполняется", "По переднему фронту", "По заднему фронту", "По любому изменению"]
    par1 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=statuses,
        #alias= f"DropDown_{i}",
        #instruction_text=f"Выберите ",
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par2 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=statuses,
    )
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par3 = hdr_cells[3].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par3,
        choices=statuses,
    )
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par4 = hdr_cells[4].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par4,
        choices=statuses,
    )
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par5 = hdr_cells[5].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par5,
        choices=statuses,
    )
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,6):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(0, 5))

    table.cell(2, 0).merge(table.cell(2, 4))
    table.cell(4, 0).merge(table.cell(4, 4))

    for row in table.rows:
        for idx, width in enumerate(table_mtrx_outs):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 

####################################################################################
############################ КОНЕЦ ТАБЛИЦА ДЛЯ МАТРИЦЫ ВЫХОДНЫХ РЕЛЕ ###############
####################################################################################


####################################################################################
############################ ТАБЛИЦА ДЛЯ МАТРИЦЫ ДИСКРЕТНЫХ ВХОДОВ ###############
####################################################################################

table_mtrx_ins = (Inches(2), Inches(4))

def add_table_mtrx_ins(doc, inputs): # новая таблица исходящих отчетов
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Дискретный вход'
    hdr_cells[1].text = 'Назначенный сигнал'

    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = r'for i in range(1, input_plate.num_of_inputs|int+1)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Дискретный вход '+'{{ loop.index }}'


    #choices_start = ["Не выполняется", "По переднему фронту", "По заднему фронту", "По любому изменению"]
    par1 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=inputs,
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(3, 0).merge(table.cell(3, 1))

    for row in table.rows:
        for idx, width in enumerate(table_mtrx_ins):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 

####################################################################################
############################ КОНЕЦ ТАБЛИЦА ДЛЯ МАТРИЦЫ ДИСКРЕТНЫХ ВХОДОВ ###########
####################################################################################


####################################################################################
############################ ТАБЛИЦА ДЛЯ СВЕТОДИОДОВ ###############
####################################################################################

table_leds = (Inches(2), Inches(2), Inches(4))

def add_table_leds(doc, statuses): # новая таблица исходящих отчетов
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Светодиод'
    hdr_cells[1].text = 'Режим работы'
    hdr_cells[2].text = 'Назначенный сигнал'    

    for i in range(0,3):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = f'for i in range(1, 17)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Светодиод '+'{{ loop.index }}'

    choices = ["С фиксацией", "Без фиксации"]
    par2 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=choices,
        default='По умолчанию'
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    par1 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=statuses,
    )
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,3):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(3, 0).merge(table.cell(3, 2))

    for row in table.rows:
        for idx, width in enumerate(table_leds):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 

####################################################################################
############################ КОНЕЦ ТАБЛИЦА ДЛЯ СВЕТОДИОДОВ ###############
####################################################################################

####################################################################################
############################ ТАБЛИЦА ДЛЯ ФУНКЦИОНАЛЬНЫХ КЛАВИШ ###############
####################################################################################

table_fks = (Inches(2), Inches(4))

def add_table_fks(doc, choices_start): # новая таблица исходящих отчетов
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Функциональная клавиша'
    hdr_cells[1].text = 'Назначенный сигнал'

    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = f'for i in range(1, 17)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Функциональная клавиша '+'{{ loop.index }}'

    par1 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=choices_start,
    )
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,2):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 1))
    table.cell(3, 0).merge(table.cell(3, 1))

    for row in table.rows:
        for idx, width in enumerate(table_fks):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 


####################################################################################
############################ ТАБЛИЦА ДЛЯ СВЕТОДИОДОВ УСОВЕРШЕНСТВОВАННАЯ ###############
####################################################################################

table_leds_new = (Inches(1.7), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5), Inches(1.5))

def add_table_leds_new(doc, statuses): # новая таблица исходящих отчетов
    table = doc.add_table(rows=5, cols=7)
    table.style = 'Стиль6'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Светодиод'
    hdr_cells[1].text = 'Режим работы'
    hdr_cells[2].text = 'Назначенный сигнал 1'    
    hdr_cells[3].text = 'Назначенный сигнал 2' 
    hdr_cells[4].text = 'Назначенный сигнал 3' 
    hdr_cells[5].text = 'Назначенный сигнал 4' 
    hdr_cells[6].text = 'Назначенный сигнал 5' 

    for i in range(0,7):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    hdr_cells = table.rows[1].cells
    tag = f'for i in range(1, 17)'
    hdr_cells[0].text = '{%tr '+ tag + ' %}'

    # четвертая строка со служебными тегами
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Светодиод '+'{{ loop.index }}' + ' (красный)'
    # четвертая строка со служебными тегами
    hdr_cells_row2 = table.rows[3].cells
    hdr_cells_row2[0].text = 'Светодиод '+'{{ loop.index }}'  + ' (зеленый)'

    choices = ["С фиксацией", "Без фиксации"]
    par2 = hdr_cells[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par2,
        choices=choices,
        default='По умолчанию')
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par1 = hdr_cells[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par1,
        choices=statuses,)
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    par21 = hdr_cells_row2[1].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par21,
        choices=choices,
        default='По умолчанию')
    hdr_cells_row2[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    par12 = hdr_cells_row2[2].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par12,
        choices=statuses,)
    hdr_cells_row2[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # первая строка
    par22 = hdr_cells[3].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par22,
        choices=statuses,)
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    par23 = hdr_cells[4].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par23,
        choices=statuses,)
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    par24 = hdr_cells[5].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par24,
        choices=statuses,)
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    par25 = hdr_cells[6].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par25,
        choices=statuses,)
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 


    # вторая строка заполнение со второго сигнала
    par13 = hdr_cells_row2[3].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par13,
        choices=statuses,)
    hdr_cells_row2[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par14 = hdr_cells_row2[4].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par14,
        choices=statuses,)
    hdr_cells_row2[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par15 = hdr_cells_row2[5].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par15,
        choices=statuses,)
    hdr_cells_row2[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    par16 = hdr_cells_row2[6].paragraphs[0]
    add_formatted_dropdown2(
        paragraph=par16,
        choices=statuses,)
    hdr_cells_row2[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,3):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек

    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(4, 0).merge(table.cell(4, 2))

    for row in table.rows:
        for idx, width in enumerate(table_leds_new):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить
    return table 


####################################################################################
################################ ТАБЛИЦА ДЛЯ ДИСКРЕТНЫХ ВХОДОВ ВЫХОДОВ ##############################
####################################################################################

table_binaries = (Inches(0.28), Inches(1.23), Inches(0.9), Inches(0.5), Inches(1.5), Inches(0.55), Inches(0.45), Inches(0.9), Inches(1.05))  #задаем ширину столбцов таблицы вывода репортов

def add_table_binaries(doc):
    table = doc.add_table(rows=5, cols=9)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Описание'
    hdr_cells[2].text = 'Наименование'
    hdr_cells[4].text = 'Значение / Диапазон'
    hdr_cells[5].text = 'Ед. изм.'
    hdr_cells[6].text = 'Шаг'
    hdr_cells[7].text = 'Значение по умолчанию'
    hdr_cells[8].text = 'Уставка'
    for i in range(0,9):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # p.runs[0].font.size = Pt(10)

    hdr_cells = table.rows[1].cells # вторая строка заголовка таблицы
    hdr_cells[2].text = 'ПО'
    hdr_cells[3].text = 'ФСУ'
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # третья строка со служебными тегами
    hdr_cells = table.rows[2].cells
    #hdr_cells[2].text = '{%tr for param_name, param_data in input_value.properties.items() %}'
    tag = f'for row in items'
    hdr_cells[2].text = '{%tr '+ tag + ' %}'
    # четвертая строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{{ loop.index }}'
    hdr_cells[1].text = '{{ row["Описание"] }}'
    hdr_cells[2].text = '{{ row["Наименование ПО"] }}'
    hdr_cells[3].text = '{{ row["Наименование ФСУ"] }}'    
    hdr_cells[4].text = '{{ row["Значение / Диапазон"]  }}'
    hdr_cells[5].text = '{{ row["Ед. изм."] }}'
    hdr_cells[6].text = '{{ row["Шаг"] }}'
    hdr_cells[7].text = '{{ row["Значение по умолчанию"] }}'
    hdr_cells[8].text = '' #'{{ param_data.setpoint }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # пятая строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}'

    set_repeat_table_header(table.rows[1])  # повторение заголовка на след странице
    for i in range(0,9):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        #set_cell_border(hdr_cells[i], bottom={"val": "double"}) # подчеркиваем заголовок двойной чертой

    # формируем финальный заголок слияниями ячеек
    table.cell(0, 2).merge(table.cell(0, 3))
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    table.cell(0, 4).merge(table.cell(1, 4))
    table.cell(0, 5).merge(table.cell(1, 5))
    table.cell(0, 6).merge(table.cell(1, 6))
    table.cell(0, 7).merge(table.cell(1, 7))
    table.cell(0, 8).merge(table.cell(1, 8))

    table.cell(2, 0).merge(table.cell(2, 8))
    table.cell(4, 0).merge(table.cell(4, 8))

    for row in table.rows:
        for idx, width in enumerate(table_binaries):
            row.cells[idx].width = width
    #add_row_table_reports(table, ('','','','','','')) # добавляем пустую строчку, чтобы двойное подчеркивание сохранить

        # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Устанавливаем размер шрифта 12 пунктов

    return table

####################################################################################
################################ КОНЕЦ ТАБЛИЦА ДЛЯ ДИСКРЕТНЫХ ФХОДОВ ВЫХОДОВ #########################
####################################################################################






####################################################################################
################################ СУММАРНАЯ ТАБЛИЦА СИГНАЛОВ #######################################

table_summ = (Inches(4), Inches(2), Inches(1), Inches(1), Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.5), Inches(0.5))  #задаем ширину столбцов таблицы

def add_summ_table(doc):
    table = doc.add_table(rows=17, cols=9)
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False

    # Устанавливаем фиксированный макет таблицы с правильным пространством имен
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Полное наименование сигналов'
    hdr_cells[1].text = 'Наименование сигналов на ФСУ'
    hdr_cells[2].text = 'Дискретные входы'
    hdr_cells[3].text = 'Выходные реле'
    hdr_cells[4].text = 'Светодиоды'
    hdr_cells[5].text = 'ФК'
    hdr_cells[6].text = 'РС'
    hdr_cells[7].text = 'РАС'
    hdr_cells[8].text = 'Пуск РАС'
    for i in range(0,9):
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Устанавливаем серый фон для всех ячеек заголовка
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        # Создаем НОВЫЙ элемент shading для каждой ячейки
        shading_elm = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>'
        )
        # Удаляем старый shading, если он есть (чтобы избежать дублирования)
        for shd in tcPr.xpath(".//w:shd"):
            tcPr.remove(shd)
        tcPr.append(shading_elm)

    set_repeat_table_header(table.rows[0]) # повторение заголовка на след странице

########################################################################################
######################## ВИРТУАЛЬНЫЕ КНОПКИ #########################################

    # третья строка со служебными тегами
    hdr_cells = table.rows[1].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_buttons() %}Виртуальные кнопки{% endif %}'
    paragraph = hdr_cells[1].paragraphs[0]
    # Делаем текст полужирным
    for run in paragraph.runs:
        run.bold = True
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    hdr_cells = table.rows[2].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_buttons() %}{%tr for row in fsu.get_fsu_buttons() %}'    
    # строка со служебными тегами
    hdr_cells = table.rows[3].cells
    hdr_cells[0].text = '{{ row["Полное наименование сигнала"] }}'
    hdr_cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
    hdr_cells[2].text = '{{ row["Дискретные входы"] }}'
    hdr_cells[3].text = '{{ row["Выходные реле"] }}'    
    hdr_cells[4].text = '{{ row["Светодиоды"]  }}'
    hdr_cells[5].text = '{{ row["ФК"] }}'
    hdr_cells[6].text = '{{ row["РС"] }}'
    hdr_cells[7].text = '{{ row["РАС"] }}'
    hdr_cells[8].text = '{{ row["Пуск РАС"] }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[8].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    # строка со служебными тегами
    hdr_cells = table.rows[4].cells
    hdr_cells[0].text = '{%tr endfor %}{% endif %}'

    table.cell(1, 0).merge(table.cell(1, 8))
    table.cell(2, 0).merge(table.cell(2, 8))
    table.cell(4, 0).merge(table.cell(4, 8))

########################################################################
#####################################################################

########################################################################################
######################## ВИРТУАЛЬНЫЕ КЛЮЧИ #########################################

    # строка со служебными тегами
    hdr_cells = table.rows[5].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_switches() %}Виртуальные ключи{% endif %}'
    paragraph = hdr_cells[1].paragraphs[0]
    # Делаем текст полужирным
    for run in paragraph.runs:
        run.bold = True
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    hdr_cells = table.rows[6].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_switches() %}{%tr for row in fsu.get_fsu_switches() %}'    
    # строка со служебными тегами
    hdr_cells = table.rows[7].cells
    hdr_cells[0].text = '{{ row["Полное наименование сигнала"] }}'
    hdr_cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
    hdr_cells[2].text = '{{ row["Дискретные входы"] }}'
    hdr_cells[3].text = '{{ row["Выходные реле"] }}'    
    hdr_cells[4].text = '{{ row["Светодиоды"]  }}'
    hdr_cells[5].text = '{{ row["ФК"] }}'
    hdr_cells[6].text = '{{ row["РС"] }}'
    hdr_cells[7].text = '{{ row["РАС"] }}'
    hdr_cells[8].text = '{{ row["Пуск РАС"] }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[8].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    # строка со служебными тегами
    hdr_cells = table.rows[8].cells
    hdr_cells[0].text = '{%tr endfor %}{% endif %}'

    table.cell(5, 0).merge(table.cell(5, 8))
    table.cell(6, 0).merge(table.cell(6, 8))
    table.cell(8, 0).merge(table.cell(8, 8))

########################################################################
#####################################################################

########################################################################################
######################## ОБЩИЕ СИГНАЛЫ #########################################

    # строка со служебными тегами
    hdr_cells = table.rows[9].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_statuses_sorted() %}Общие сигналы функциональной логики{% endif %}'
    paragraph = hdr_cells[1].paragraphs[0]
    # Делаем текст полужирным
    for run in paragraph.runs:
        run.bold = True    
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    hdr_cells = table.rows[10].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_statuses_sorted() %}{%tr for row in fsu.get_fsu_statuses_sorted() %}'    
    # строка со служебными тегами
    hdr_cells = table.rows[11].cells
    hdr_cells[0].text = '{{ row["Полное наименование сигнала"] }}'
    hdr_cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
    hdr_cells[2].text = '{{ row["Дискретные входы"] }}'
    hdr_cells[3].text = '{{ row["Выходные реле"] }}'    
    hdr_cells[4].text = '{{ row["Светодиоды"]  }}'
    hdr_cells[5].text = '{{ row["ФК"] }}'
    hdr_cells[6].text = '{{ row["РС"] }}'
    hdr_cells[7].text = '{{ row["РАС"] }}'
    hdr_cells[8].text = '{{ row["Пуск РАС"] }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[8].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    # строка со служебными тегами
    hdr_cells = table.rows[12].cells
    hdr_cells[0].text = '{%tr endfor %}{% endif %}'

    table.cell(9, 0).merge(table.cell(9, 8))
    table.cell(10, 0).merge(table.cell(10, 8))
    table.cell(12, 0).merge(table.cell(12, 8))

################################################################################
############################################################################

########################################################################################
######################## СИСТЕМНЫЕ СИГНАЛЫ #########################################

    # строка со служебными тегами
    hdr_cells = table.rows[13].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_sys_statuses_sorted() %}Системные сигналы{% endif %}'
    paragraph = hdr_cells[1].paragraphs[0]
    # Делаем текст полужирным
    for run in paragraph.runs:
        run.bold = True    
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   
    hdr_cells = table.rows[14].cells
    hdr_cells[1].text = r'{% if fsu.get_fsu_sys_statuses_sorted() %}{%tr for row in fsu.get_fsu_sys_statuses_sorted() %}'    
    # строка со служебными тегами
    hdr_cells = table.rows[15].cells
    hdr_cells[0].text = '{{ row["Полное наименование сигнала"] }}'
    hdr_cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
    hdr_cells[2].text = '{{ row["Дискретные входы"] }}'
    hdr_cells[3].text = '{{ row["Выходные реле"] }}'    
    hdr_cells[4].text = '{{ row["Светодиоды"]  }}'
    hdr_cells[5].text = '{{ row["ФК"] }}'
    hdr_cells[6].text = '{{ row["РС"] }}'
    hdr_cells[7].text = '{{ row["РАС"] }}'
    hdr_cells[8].text = '{{ row["Пуск РАС"] }}'

    hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[4].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    hdr_cells[5].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[6].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[7].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_cells[8].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER    
    # строка со служебными тегами
    hdr_cells = table.rows[16].cells
    hdr_cells[0].text = '{%tr endfor %}{% endif %}'

    table.cell(13, 0).merge(table.cell(13, 8))
    table.cell(14, 0).merge(table.cell(14, 8))
    table.cell(16, 0).merge(table.cell(16, 8))

################################################################################
############################################################################



    # Приведение оформления таблицы
    for row in table.rows:
        for idx, width in enumerate(table_summ):
            row.cells[idx].width = width
        # Устанавливаем высоту шрифта (11 пунктов) для всех ячеек таблицы
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)  # Устанавливаем размер шрифта 12 пунктов

    return table

####################################################################################
################################ КОНЕЦ СУММАРНАЯ ТАБЛИЦА СИГНАЛОВ #########################
####################################################################################



####################################################################################
################################ СУММАРНАЯ ТАБЛИЦА СИГНАЛОВ ВЕРСИЯ 2 #######################################
def add_summ_table2(doc, isVirtKey=False, isVirtSwitch=False, isStatuses=False, isSysStatuses=False):
    table = doc.add_table(rows=1, cols=9)  # Начинаем с одной строки
    table.style = 'Сетка таблицы51'
    table.allow_autofit = False
    table._tbl.xpath('./w:tblPr')[0].append(
        parse_xml(r'<w:tblLayout xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="fixed"/>')
    )

    # Добавляем заголовок
    hdr_cells = table.rows[0].cells
    for i, text in enumerate([
        'Полное наименование сигналов', 'Наименование сигналов на ФСУ',
        'Дискретные входы', 'Выходные реле', 'Светодиоды', 'ФК', 'РС', 'РАС', 'Пуск РАС'
    ]):
        hdr_cells[i].text = text
        p = hdr_cells[i].paragraphs[0]
        p.style = 'ДОК Таблица Заголовок'
        set_cell_vertical_alignment(hdr_cells[i], align="center")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
        for shd in tcPr.xpath(".//w:shd"):
            tcPr.remove(shd)
        tcPr.append(shading_elm)

    set_repeat_table_header(table.rows[0])

    current_row = 1  # Следующая строка после заголовка

    # --- Виртуальные кнопки ---
    if isVirtKey:
        table.add_row()
        row = table.rows[current_row]
        cell = row.cells[0]
        cell.text = r'{% if fsu.get_fsu_buttons() %}Виртуальные кнопки{% endif %}'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[1].text = r'{% if fsu.get_fsu_buttons() %}{%tr for row in fsu.get_fsu_buttons() %}'
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{{ row["Полное наименование сигнала"] }}'
        row.cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
        row.cells[2].text = '{{ row["Дискретные входы"] }}'
        row.cells[3].text = '{{ row["Выходные реле"] }}'
        row.cells[4].text = '{{ row["Светодиоды"] }}'
        row.cells[5].text = '{{ row["ФК"] }}'
        row.cells[6].text = '{{ row["РС"] }}'
        row.cells[7].text = '{{ row["РАС"] }}'
        row.cells[8].text = '{{ row["Пуск РАС"] }}'
        for i in range(9):
            row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{%tr endfor %}{% endif %}'
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

    # --- Виртуальные ключи ---
    if isVirtSwitch:
        table.add_row()
        row = table.rows[current_row]
        cell = row.cells[0]
        cell.text = r'{% if fsu.get_fsu_switches() %}Виртуальные ключи{% endif %}'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[1].text = r'{% if fsu.get_fsu_switches() %}{%tr for row in fsu.get_fsu_switches() %}'
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{{ row["Полное наименование сигнала"] }}'
        row.cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
        row.cells[2].text = '{{ row["Дискретные входы"] }}'
        row.cells[3].text = '{{ row["Выходные реле"] }}'
        row.cells[4].text = '{{ row["Светодиоды"] }}'
        row.cells[5].text = '{{ row["ФК"] }}'
        row.cells[6].text = '{{ row["РС"] }}'
        row.cells[7].text = '{{ row["РАС"] }}'
        row.cells[8].text = '{{ row["Пуск РАС"] }}'
        for i in range(9):
            row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{%tr endfor %}{% endif %}'
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

    # --- Общие сигналы ---
    if isStatuses:
        table.add_row()
        row = table.rows[current_row]
        cell = row.cells[0]
        cell.text = r'{% if fsu.get_fsu_statuses_sorted() %}Общие сигналы функциональной логики{% endif %}'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[1].text = r'{% if fsu.get_fsu_statuses_sorted() %}{%tr for row in fsu.get_fsu_statuses_sorted() %}'
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{{ row["Полное наименование сигнала"] }}'
        row.cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
        row.cells[2].text = '{{ row["Дискретные входы"] }}'
        row.cells[3].text = '{{ row["Выходные реле"] }}'
        row.cells[4].text = '{{ row["Светодиоды"] }}'
        row.cells[5].text = '{{ row["ФК"] }}'
        row.cells[6].text = '{{ row["РС"] }}'
        row.cells[7].text = '{{ row["РАС"] }}'
        row.cells[8].text = '{{ row["Пуск РАС"] }}'
        for i in range(9):
            row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{%tr endfor %}{% endif %}'
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

    # --- Системные сигналы ---
    if isSysStatuses:
        table.add_row()
        row = table.rows[current_row]
        cell = row.cells[0]
        cell.text = r'{% if fsu.get_fsu_sys_statuses_sorted() %}Системные сигналы{% endif %}'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[1].text = r'{% if fsu.get_fsu_sys_statuses_sorted() %}{%tr for row in fsu.get_fsu_sys_statuses_sorted() %}'
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{{ row["Полное наименование сигнала"] }}'
        row.cells[1].text = '{{ row["Наименование сигналов на ФСУ"] }}'
        row.cells[2].text = '{{ row["Дискретные входы"] }}'
        row.cells[3].text = '{{ row["Выходные реле"] }}'
        row.cells[4].text = '{{ row["Светодиоды"] }}'
        row.cells[5].text = '{{ row["ФК"] }}'
        row.cells[6].text = '{{ row["РС"] }}'
        row.cells[7].text = '{{ row["РАС"] }}'
        row.cells[8].text = '{{ row["Пуск РАС"] }}'
        for i in range(9):
            row.cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        current_row += 1

        table.add_row()
        row = table.rows[current_row]
        row.cells[0].text = '{%tr endfor %}{% endif %}'
        table.cell(current_row, 0).merge(table.cell(current_row, 8))
        current_row += 1

    # --- Форматирование таблицы ---
    for row in table.rows:
        for idx, width in enumerate(table_summ):
            row.cells[idx].width = width
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    return table